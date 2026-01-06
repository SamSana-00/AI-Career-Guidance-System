from flask import Flask, request, redirect, url_for, session, render_template, flash, send_file, jsonify
import PyPDF2
import ollama
import hashlib
from flask_session import Session
from docx import Document
import io
import re
from markdown import markdown
import logging
import json
from datetime import datetime
import pandas as pd
import numpy as np
from functools import lru_cache
import time
import webbrowser
import threading
import json
import os

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = "your_secure_secret_key"  # Change this in production
app.config["SESSION_TYPE"] = "filesystem"
Session(app)

def load_users():
    """Load users from JSON file"""
    try:
        if os.path.exists('users.json'):
            with open('users.json', 'r') as file:
                return json.load(file)
        return {}
    except Exception as e:
        logger.error(f"Error loading users: {e}")
        return {}

def save_users():
    """Save users to JSON file"""
    try:
        with open('users.json', 'w') as file:
            json.dump(users, file)
    except Exception as e:
        logger.error(f"Error saving users: {e}")

# Initialize users from file
users = load_users()

def hash_password(password):
    """Hash password using SHA-256"""
    return hashlib.sha256(password.encode()).hexdigest()

def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF with error handling"""
    try:
        reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            try:
                text += page.extract_text() + "\n"
            except Exception as e:
                logger.error(f"Error extracting text from page: {e}")
        return text
    except Exception as e:
        logger.error(f"Error reading PDF: {e}")
        return ""

def extract_keywords_from_pdf(pdf_content):
    """Extract relevant keywords and information from PDF content"""
    try:
        # Remove common words and get key phrases
        common_words = set(['and', 'the', 'is', 'at', 'which', 'on', 'the', 'a', 'an', 'of', 'for', 'to', 'in'])
        
        # Extract sections using regex
        education = re.findall(r'education:(.+?)(?:\n\n|\Z)', pdf_content, re.DOTALL | re.IGNORECASE)
        skills = re.findall(r'skills:(.+?)(?:\n\n|\Z)', pdf_content, re.DOTALL | re.IGNORECASE)
        experience = re.findall(r'experience:(.+?)(?:\n\n|\Z)', pdf_content, re.DOTALL | re.IGNORECASE)
        
        # Extract technical terms
        technical_terms = re.findall(r'\b[A-Za-z\+\#]+(?:\.[A-Za-z\+\#]+)*\b', pdf_content)
        technical_terms = [term for term in technical_terms if term.lower() not in common_words]
        
        relevant_info = f"""
        Key Information from Resume:
        Education Background: {' '.join(education[:300]) if education else 'Not specified'}
        Skills Highlighted: {' '.join(skills[:300]) if skills else 'Not specified'}
        Experience Summary: {' '.join(experience[:300]) if experience else 'Not specified'}
        Technical Terms: {', '.join(set(technical_terms[:50]))}
        """
        return relevant_info
    except Exception as e:
        logger.error(f"Error extracting keywords: {e}")
        return "Error processing document content"

@lru_cache(maxsize=100)
# First, make sure Ollama is properly installed and running
# Then, modify the generate_career_guidance function to better handle errors:

def generate_career_guidance(profile_str, pdf_content):
    """Generate personalized career guidance using LLM with caching"""
    try:
        # Add debug logging
        logger.info(f"Starting career guidance generation with profile length: {len(profile_str)} and PDF content length: {len(pdf_content)}")
        
        relevant_keywords = extract_keywords_from_pdf(pdf_content)
        logger.info(f"Extracted keywords successfully, length: {len(relevant_keywords)}")
        
        prompt = f"""
        You are a career guidance expert. Based on the following student profile and their resume content, provide highly personalized career guidance.
        
        Student Profile:
        {profile_str}
        
        Key points from their resume/documents:
        {relevant_keywords}

        Provide detailed, personalized career guidance addressing:

        1. Career Paths
        Analyze their profile and suggest 3-4 specific career paths that align with their:
        - Subject preferences
        - Learning goals
        - Current skills (from resume)
        
        2. Required Skills & Certifications
        For each career path mentioned above, list:
        - Technical skills needed
        - Soft skills required
        - Specific certifications or degrees recommended
        
        3. Job Market Analysis
        For each career path:
        - Current job market demand
        - Expected salary ranges (entry-level to experienced)
        - Growth potential over next 5 years
        
        4. Personalized Learning Roadmap
        Create a structured, quarter-by-quarter learning plan for the next 2 years:
        Q1: [Specific goals and activities]
        Q2: [Specific goals and activities]
        Q3: [Specific goals and activities]
        Q4: [Specific goals and activities]
        Q5: [Specific goals and activities]
        Q6: [Specific goals and activities]
        Q7: [Specific goals and activities]
        Q8: [Specific goals and activities]

        Format the response using proper markdown with clear headings and bullet points.
        Ensure all suggestions are specifically tailored to their profile and resume content.
        """

        # Check if Ollama is running and accessible
        try:
            start_time = time.time()
            logger.info("Attempting to connect to Ollama...")
            
            # Verify Ollama model is available - add a simple test call with timeout
            test_response = ollama.chat(
                model="deepseek-r1:1.5b",
                messages=[{"role": "user", "content": "test"}],
                options={"timeout": 5}  # Add timeout option
            )
            logger.info("Ollama test connection successful")
            
            # Now send the actual request
            response = ollama.chat(
                model="deepseek-r1:1.5b", 
                messages=[{"role": "user", "content": prompt}]
            )
            end_time = time.time()
            
            logger.info(f"LLM response generated in {end_time - start_time:.2f} seconds")
            
            # Validate the response structure
            if 'message' not in response or 'content' not in response['message']:
                logger.error(f"Unexpected response format from Ollama: {response}")
                return "Error: Received unexpected response format from the language model."
                
            return clean_llm_output(response['message']['content'])
            
        except Exception as e:
            logger.error(f"Ollama communication error: {e}")
            # Fallback response if Ollama fails
            return """
            # Career Guidance Report

            ## Technical Issues

            I apologize, but we're experiencing technical difficulties with our AI career guidance system. Here are some alternatives:

            1. Please try again later when our systems are back online
            2. Contact support for assistance
            3. Try uploading a different PDF format

            Our team has been notified of this issue and is working to resolve it as quickly as possible.
            """
            
    except Exception as e:
        logger.error(f"Error generating career guidance: {e}", exc_info=True)  # Add full traceback
        return "Error generating career guidance. Please try again."

def clean_llm_output(response_text):
    """Remove unnecessary tags and clean response"""
    try:
        response_text = re.sub(r"<think>.*?</think>", "", response_text, flags=re.DOTALL)
        response_text = re.sub(r"<.*?>", "", response_text, flags=re.DOTALL)
        response_text = re.sub(r'\n{3,}', '\n\n', response_text)
        return response_text.strip()
    except Exception as e:
        logger.error(f"Error cleaning LLM output: {e}")
        return response_text

def create_doc_file(content):
    """Convert markdown content to properly formatted .docx file"""
    try:
        doc = Document()
        doc.add_heading("Career Guidance Report", level=0)
        
        sections = content.split('\n#')
        
        for section in sections:
            if not section.strip():
                continue
            
            lines = section.split('\n')
            title = lines[0].strip('# ')
            doc.add_heading(title, level=1)
            
            current_list = []
            for line in lines[1:]:
                if not line.strip():
                    if current_list:
                        for item in current_list:
                            doc.add_paragraph(item, style='List Bullet')
                        current_list = []
                    continue
                
                if line.startswith('##'):
                    if current_list:
                        for item in current_list:
                            doc.add_paragraph(item, style='List Bullet')
                        current_list = []
                    doc.add_heading(line.strip('# '), level=2)
                elif line.startswith('- '):
                    current_list.append(line.strip('- '))
                elif line.startswith('Q'):
                    if current_list:
                        for item in current_list:
                            doc.add_paragraph(item, style='List Bullet')
                        current_list = []
                    doc.add_paragraph(line, style='Intense Quote')
                else:
                    if current_list:
                        for item in current_list:
                            doc.add_paragraph(item, style='List Bullet')
                        current_list = []
                    doc.add_paragraph(line)
        
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
    except Exception as e:
        logger.error(f"Error creating document: {e}")
        raise

def generate_skill_gap_analysis(target_role, current_skills, resume_content):
    """Generate skill gap analysis using LLM"""
    try:
        # Log the request details
        logger.info(f"Generating skill gap analysis for role: {target_role}")
        logger.info(f"Current skills: {current_skills}")
        logger.info(f"Resume content length: {len(resume_content) if resume_content else 0}")
        
        # Ensure current_skills is a list
        if current_skills is None:
            current_skills = []
        elif isinstance(current_skills, str):
            current_skills = [current_skills]
        
        current_skills_str = ', '.join(current_skills) if current_skills else 'None provided'
        
        prompt = f"""
        As a career development expert, analyze the skill gap for the following:
        
        Target Role: {target_role}
        Current Skills: {current_skills_str}
        Resume Content: {resume_content[:500] if resume_content else 'None provided'}

        Please analyze and provide:
        1. All required skills for the target role (be comprehensive)
        2. Skills that are missing (skills required for the role but not in current skills)
        3. Skill proficiency recommendations
        4. Development priorities (ordered by importance)
        
        Current skills to check against: {current_skills}
        
        IMPORTANT: Compare the required skills with the current skills list above and identify what's missing.
        
        Format your response EXACTLY as valid JSON with these keys:
        {{
            "required_skills": ["list of ALL skills needed for {target_role}"],
            "missing_skills": ["skills from required_skills that are NOT in current_skills"],
            "proficiency_recommendations": {{
                "skill_name": "proficiency level and recommendation"
            }},
            "development_priorities": ["ordered list of most important skills to develop first"]
        }}
        
        Example for a Data Scientist role:
        {{
            "required_skills": ["Python", "Machine Learning", "Statistics", "SQL", "Data Visualization", "Deep Learning", "R", "Tableau"],
            "missing_skills": ["Machine Learning", "Deep Learning", "R", "Tableau"],
            "proficiency_recommendations": {{
                "Python": "Advanced level needed for data manipulation",
                "Machine Learning": "Intermediate to advanced level required",
                "Statistics": "Strong foundation in descriptive and inferential statistics"
            }},
            "development_priorities": ["Machine Learning", "Statistics", "Deep Learning", "Data Visualization"]
        }}
        
        Ensure your output is ONLY the JSON object, with no additional text before or after.
        """
        
        # Test connection to Ollama first
        logger.info("Testing connection to Ollama")
        test_response = ollama.chat(
            model="deepseek-r1:1.5b",
            messages=[{"role": "user", "content": "test"}],
            options={"timeout": 5}
        )
        logger.info("Ollama test connection successful")
        
        # Send the actual request
        response = ollama.chat(
            model="deepseek-r1:1.5b",
            messages=[{"role": "user", "content": prompt}]
        )
        
        # Extract and log response content
        response_content = response['message']['content']
        logger.info(f"Received response from Ollama, length: {len(response_content)}")
        logger.info(f"Response preview: {response_content[:200]}...")
        
        # Clean and parse JSON
        cleaned_content = response_content.strip()
        
        # Remove any thinking tags or markdown formatting
        import re
        cleaned_content = re.sub(r'<think>.*?</think>', '', cleaned_content, flags=re.DOTALL)
        cleaned_content = re.sub(r'```json\s*', '', cleaned_content)
        cleaned_content = re.sub(r'```\s*', '', cleaned_content)
        cleaned_content = cleaned_content.strip()
        
        logger.info(f"Cleaned content for JSON parsing: {cleaned_content[:200]}...")
        
        # Try to parse JSON
        try:
            result = json.loads(cleaned_content)
            logger.info("Successfully parsed JSON response")
            
            # Validate and fix the response
            required_keys = ['required_skills', 'missing_skills', 'proficiency_recommendations', 'development_priorities']
            for key in required_keys:
                if key not in result:
                    logger.warning(f"Missing expected key in response: {key}")
                    if key == 'proficiency_recommendations':
                        result[key] = {}
                    else:
                        result[key] = []
            
            # Double-check missing skills calculation
            if result.get('required_skills') and current_skills:
                # Convert to lowercase for comparison
                current_skills_lower = [skill.lower().strip() for skill in current_skills]
                required_skills_lower = [skill.lower().strip() for skill in result['required_skills']]
                
                # Find missing skills
                actual_missing = []
                for req_skill in result['required_skills']:
                    req_skill_lower = req_skill.lower().strip()
                    if not any(req_skill_lower in curr_skill or curr_skill in req_skill_lower 
                             for curr_skill in current_skills_lower):
                        actual_missing.append(req_skill)
                
                result['missing_skills'] = actual_missing
                logger.info(f"Recalculated missing skills: {actual_missing}")
            
            # If still no missing skills but we have required skills and current skills don't match
            if not result.get('missing_skills') and result.get('required_skills') and current_skills:
                # Fallback: assume some skills are missing
                required = result['required_skills']
                if len(required) > len(current_skills):
                    result['missing_skills'] = required[len(current_skills):]
            
            logger.info(f"Final analysis result: {result}")
            return result
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON: {e}")
            logger.error(f"Content that failed to parse: {cleaned_content}")
            
            # Return intelligent fallback based on role
            fallback_skills = get_fallback_skills_for_role(target_role)
            missing_skills = [skill for skill in fallback_skills if skill not in current_skills]
            
            return {
                "required_skills": fallback_skills,
                "missing_skills": missing_skills,
                "proficiency_recommendations": {skill: f"Required for {target_role}" for skill in fallback_skills[:3]},
                "development_priorities": missing_skills[:4] if missing_skills else fallback_skills[:4]
            }
        
    except Exception as e:
        logger.error(f"Error generating skill gap analysis: {e}", exc_info=True)
        
        # Return intelligent fallback based on role
        fallback_skills = get_fallback_skills_for_role(target_role)
        missing_skills = [skill for skill in fallback_skills if skill not in current_skills]
        
        return {
            "required_skills": fallback_skills,
            "missing_skills": missing_skills,
            "proficiency_recommendations": {skill: f"Required for {target_role}" for skill in fallback_skills[:3]},
            "development_priorities": missing_skills[:4] if missing_skills else fallback_skills[:4]
        }

def get_fallback_skills_for_role(target_role):
    """Get fallback skills based on role keywords"""
    role_lower = target_role.lower()
    
    # Define common skills for different role types
    skill_mapping = {
        'data scientist': ['Python', 'Machine Learning', 'Statistics', 'SQL', 'Data Visualization', 'Pandas', 'Scikit-learn'],
        'software engineer': ['Programming', 'Algorithms', 'Data Structures', 'Version Control (Git)', 'Problem Solving', 'Testing'],
        'web developer': ['HTML/CSS', 'JavaScript', 'React or Vue.js', 'Node.js', 'Database Management', 'Responsive Design'],
        'machine learning': ['Python', 'Machine Learning', 'Deep Learning', 'TensorFlow/PyTorch', 'Statistics', 'Data Preprocessing'],
        'product manager': ['Product Strategy', 'Market Research', 'Agile/Scrum', 'Data Analysis', 'Communication', 'Roadmap Planning'],
        'ux designer': ['User Research', 'Wireframing', 'Prototyping', 'Figma/Sketch', 'Usability Testing', 'Information Architecture'],
        'devops': ['Docker', 'Kubernetes', 'CI/CD', 'AWS/Azure', 'Linux', 'Infrastructure as Code', 'Monitoring'],
        'analyst': ['Excel', 'SQL', 'Data Analysis', 'Reporting', 'Statistics', 'Business Intelligence', 'Critical Thinking']
    }
    
    # Find matching skills
    for key, skills in skill_mapping.items():
        if key in role_lower:
            return skills
    
    # Default fallback for any role
    return ['Communication', 'Problem Solving', 'Critical Thinking', 'Time Management', 'Teamwork', 'Adaptability']

def recommend_learning_resources(missing_skills, profile=None):
    """Recommend learning resources for missing skills with profile context"""
    try:
        # Log the request details
        logger.info(f"Recommending learning resources for: {missing_skills}")
        
        if not missing_skills:
            logger.warning("No missing skills provided")
            return {}
        
        # Create profile context
        profile_context = ""
        if profile:
            profile_context = "\nCandidate Profile:\n"
            for key, value in profile.items():
                if isinstance(value, list):
                    profile_context += f"- {key}: {', '.join(value)}\n"
                else:
                    profile_context += f"- {key}: {value}\n"
        
        prompt = f"""
        Recommend specific learning resources for these skills:
        {', '.join(missing_skills)}
        {profile_context}
        
        For each skill provide:
        1. Online courses (with platforms)
        2. Certifications
        3. Practice projects
        4. Estimated time to learn
        
        IMPORTANT: Tailor the recommendations based on the candidate's profile information.
        
        Format your response EXACTLY as valid JSON with skill names as keys and recommendations as values.
        
        Example format:
        {{
            "Skill 1": {{
                "courses": ["Course 1 (Platform)", "Course 2 (Platform)"],
                "certifications": ["Certification 1", "Certification 2"],
                "projects": ["Project 1", "Project 2"],
                "estimated_time": "X months"
            }},
            "Skill 2": {{
                "courses": ["Course 1 (Platform)", "Course 2 (Platform)"],
                "certifications": ["Certification 1", "Certification 2"],
                "projects": ["Project 1", "Project 2"],
                "estimated_time": "X months"
            }}
        }}
        
        Ensure your output is ONLY the JSON object, with no additional text before or after.
        """
        
        # Send the request
        response = ollama.chat(
            model="deepseek-r1:1.5b",
            messages=[{"role": "user", "content": prompt}]
        )
        
        # Extract and log response content
        response_content = response['message']['content']
        logger.info(f"Received response from Ollama, length: {len(response_content)}")
        
        # Clean and parse JSON
        cleaned_content = response_content.strip()
        if cleaned_content.startswith("```json"):
            cleaned_content = cleaned_content[7:]
        if cleaned_content.startswith("```"):
            cleaned_content = cleaned_content[3:]
        if cleaned_content.endswith("```"):
            cleaned_content = cleaned_content[:-3]
        
        cleaned_content = cleaned_content.strip()
        
        try:
            result = json.loads(cleaned_content)
            logger.info("Successfully parsed JSON response for learning resources")
            return result
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON for learning resources: {e}")
            
            # Return fallback structure
            fallback = {}
            for skill in missing_skills:
                fallback[skill] = {
                    "courses": [f"Online course for {skill} (Udemy)", f"Advanced {skill} (Coursera)"],
                    "certifications": [f"{skill} Professional Certification"],
                    "projects": [f"Build a simple {skill} project", f"Contribute to open source {skill} project"],
                    "estimated_time": "3-6 months"
                }
            return fallback
        
    except Exception as e:
        logger.error(f"Error recommending learning resources: {e}", exc_info=True)
        
        # Return fallback structure on error
        fallback = {}
        for skill in missing_skills:
            fallback[skill] = {
                "courses": [f"Online course for {skill} (Udemy)", f"Advanced {skill} (Coursera)"],
                "certifications": [f"{skill} Professional Certification"],
                "projects": [f"Build a simple {skill} project", f"Contribute to open source {skill} project"],
                "estimated_time": "3-6 months"
            }
        return fallback

def generate_interview_prep(role, experience_level, profile, resume_content):
    """Generate personalized interview preparation content with improved error handling"""
    try:
        # Log the request details
        logger.info(f"Generating interview prep for role: {role}, experience level: {experience_level}")
        logger.info(f"Profile data available: {bool(profile)}")
        logger.info(f"Resume content length: {len(resume_content) if resume_content else 0}")
        
        # Create profile text for the prompt
        profile_text = ""
        if profile:
            profile_text = "\nCandidate Profile:\n"
            for key, value in profile.items():
                if isinstance(value, list):
                    profile_text += f"- {key}: {', '.join(value)}\n"
                else:
                    profile_text += f"- {key}: {value}\n"
        
        prompt = f"""
        As an expert interviewer for {role} positions, generate:
        
        1. Technical questions specific to {role}
        2. Behavioral questions relevant to {experience_level} level
        3. Sample answers based on this candidate's profile:
        {profile_text}
        
        Resume content: {resume_content[:500] if resume_content else "Not available"}
        
        Include:
        - 5 technical questions with answers
        - 5 behavioral questions with answer frameworks
        - 3 questions candidate should ask interviewer
        
        Format as valid JSON with exactly these keys:
        - technical_questions (list of dict with 'question' and 'answer' keys)
        - behavioral_questions (list of dict with 'question' and 'framework' keys)
        - questions_to_ask (list)
        
        IMPORTANT: Your response MUST be valid JSON only, with no additional text before or after the JSON object.
        """
        
        # Test connection to Ollama first
        logger.info("Testing connection to Ollama")
        try:
            test_response = ollama.chat(
                model="deepseek-r1:1.5b",
                messages=[{"role": "user", "content": "test"}],
                options={"timeout": 5}
            )
            logger.info("Ollama test connection successful")
        except Exception as e:
            logger.error(f"Ollama connection test failed: {e}")
            raise Exception("Failed to connect to Ollama. Please check if the service is running.")
        
        # Send the actual request
        response = ollama.chat(
            model="deepseek-r1:1.5b",
            messages=[{"role": "user", "content": prompt}]
        )
        
        # Extract and log response content
        response_content = response['message']['content']
        logger.info(f"Received response from Ollama, length: {len(response_content)}")
        logger.info(f"Response preview: {response_content[:100]}...")
        
        # Clean and parse JSON
        cleaned_content = response_content.strip()
        
        # Remove markdown code block indicators if present
        if cleaned_content.startswith("```json"):
            cleaned_content = cleaned_content[7:]
        elif cleaned_content.startswith("```"):
            cleaned_content = cleaned_content[3:]
        if cleaned_content.endswith("```"):
            cleaned_content = cleaned_content[:-3]
        
        cleaned_content = cleaned_content.strip()
        logger.info(f"Cleaned content for JSON parsing, length: {len(cleaned_content)}")
        logger.info(f"Attempting to parse JSON: {cleaned_content[:50]}...")
        
        try:
            result = json.loads(cleaned_content)
            logger.info("Successfully parsed JSON response")
            
            # Validate expected keys
            required_keys = ['technical_questions', 'behavioral_questions', 'questions_to_ask']
            for key in required_keys:
                if key not in result:
                    logger.warning(f"Missing expected key in response: {key}")
                    if key == 'technical_questions' or key == 'behavioral_questions':
                        result[key] = [{"question": f"Sample {key.replace('_', ' ')} question", "answer" if key == 'technical_questions' else "framework": "Sample answer"}]
                    else:
                        result[key] = ["Sample question to ask"]
            
            return result
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error: {e}")
            logger.error(f"Content that failed to parse: {cleaned_content[:200]}...")
            
            # Generate a fallback response structure
            fallback = {
                "technical_questions": [
                    {"question": f"What technical skills are most important for a {role} position?", 
                     "answer": f"For a {role} position, key technical skills include... [This is a fallback response due to a processing error]"},
                    {"question": f"How would you approach [technical challenge] in a {role} position?", 
                     "answer": "I would first analyze the requirements, then... [This is a fallback response due to a processing error]"},
                    {"question": f"What tools or technologies are you proficient in that relate to {role}?", 
                     "answer": "I'm proficient in... [This is a fallback response due to a processing error]"},
                    {"question": f"Describe your experience with [key technology] for {role}.", 
                     "answer": "My experience includes... [This is a fallback response due to a processing error]"},
                    {"question": f"How do you stay current with developments in the {role} field?", 
                     "answer": "I stay current by... [This is a fallback response due to a processing error]"}
                ],
                "behavioral_questions": [
                    {"question": "Tell me about a time you overcame a challenge in your work.", 
                     "framework": "Use the STAR method: Situation, Task, Action, Result. [This is a fallback response due to a processing error]"},
                    {"question": "How do you handle tight deadlines?", 
                     "framework": "Describe your prioritization process and how you manage stress. [This is a fallback response due to a processing error]"},
                    {"question": "Give an example of a successful team project you contributed to.", 
                     "framework": "Focus on your specific contributions and the overall team success. [This is a fallback response due to a processing error]"},
                    {"question": "How do you handle disagreements with team members?", 
                     "framework": "Emphasize communication skills and finding common ground. [This is a fallback response due to a processing error]"},
                    {"question": "Describe a situation where you had to learn something new quickly.", 
                     "framework": "Show your adaptability and learning approach. [This is a fallback response due to a processing error]"}
                ],
                "questions_to_ask": [
                    f"What does success look like for a {role} in the first 3-6 months?",
                    "How would you describe the team culture and working environment?",
                    "What are the biggest challenges facing the team right now?"
                ]
            }
            
            # Customize the fallback response based on the role
            if "data" in role.lower() or "analyst" in role.lower():
                fallback["technical_questions"][0]["question"] = "What data analysis tools are you proficient with?"
            elif "developer" in role.lower() or "engineer" in role.lower():
                fallback["technical_questions"][0]["question"] = "What programming languages are you most comfortable with?"
            elif "manager" in role.lower() or "leader" in role.lower():
                fallback["technical_questions"][0]["question"] = "How do you approach team management and delegation?"
                
            logger.info("Generated fallback interview preparation content")
            return fallback
            
    except Exception as e:
        logger.error(f"Error generating interview preparation: {e}", exc_info=True)
        
        # Return a minimal fallback structure on critical error
        return {
            "technical_questions": [
                {"question": "Sample technical question", "answer": "Sample answer"}
            ],
            "behavioral_questions": [
                {"question": "Sample behavioral question", "framework": "Sample framework"}
            ],
            "questions_to_ask": [
                "Sample question to ask"
            ]
        }

def generate_career_comparison(paths, education_costs, timeframes):
    """Generate career path comparison and ROI analysis"""
    prompt = f"""
    As a career and financial analyst, compare these career paths:
    Career Paths: {', '.join(paths)}
    Education Costs: {', '.join(education_costs)}
    Timeframes: {', '.join(timeframes)}
    
    Provide:
    1. 5-year salary projections
    2. Education ROI analysis
    3. Industry growth trends
    4. Risk factors
    
    Format as JSON with these keys:
    - salary_projections (dict with years as keys)
    - roi_analysis (dict with paths as keys)
    - industry_trends (dict with paths as keys)
    - risk_factors (dict with paths as keys)
    """
    
    response = ollama.chat(
        model="deepseek-r1:1.5b",
        messages=[{"role": "user", "content": prompt}]
    )
    
    return json.loads(response['message']['content'])

# Routes
@app.route("/dashboard")
def dashboard():
    if "user" not in session:
        return redirect(url_for("login"))
    
    user_profile = session.get("profile", {})
    recent_guidance = session.get("career_advice_markdown", "")
    interview_prep = session.get("interview_data", {})
    skill_gap_data = session.get("skill_gap_analysis", {})
    
    # Calculate completion status
    completion_stats = {
        "profile_complete": bool(user_profile and session.get("pdf_content")),
        "guidance_generated": bool(recent_guidance),
        "interview_prep_done": bool(interview_prep),
        "skill_gap_done": bool(skill_gap_data)
    }
    
    # Calculate overall progress percentage
    total_features = len(completion_stats)
    completed_features = sum(completion_stats.values())
    progress_percentage = int((completed_features / total_features) * 100)
    
    return render_template(
        "dashboard.html",
        username=session["user"],
        profile=user_profile,
        has_guidance=bool(recent_guidance),
        has_interview_prep=bool(interview_prep),
        has_skill_gap=bool(skill_gap_data),
        completion_stats=completion_stats,
        progress_percentage=progress_percentage
    )
@app.route("/", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = hash_password(request.form["password"])

        if username in users and users[username] == password:
            session["user"] = username
            return redirect(url_for("dashboard"))
        flash("Invalid username or password", "danger")

    return render_template("login.html")

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        username = request.form["username"]
        password = hash_password(request.form["password"])

        if username in users:
            flash("Username already taken", "danger")
        else:
            users[username] = password
            save_users()  # Save users to file after adding new user
            flash("Account created! Please log in.", "success")
            return redirect(url_for("login"))

    return render_template("signup.html")

@app.route("/setup", methods=["GET", "POST"])
def index():
    if "user" not in session:
        return redirect(url_for("login"))

    if request.method == "POST":
        try:
            pdf_file = request.files.get("pdf_file")
            if not pdf_file:
                flash("Please upload a PDF file", "warning")
                return redirect(url_for("index"))
            
            # Log PDF details
            logger.info(f"PDF file uploaded: {pdf_file.filename}, size: {pdf_file.content_length} bytes")
            
            # Extract and validate PDF content
            pdf_content = extract_text_from_pdf(pdf_file)
            if not pdf_content or len(pdf_content.strip()) < 10:
                logger.error(f"PDF extraction failed or produced minimal content: {pdf_content[:100]}")
                flash("Could not extract text from the PDF. Please try a different file.", "warning")
                return redirect(url_for("index"))
            
            logger.info(f"Successfully extracted {len(pdf_content)} characters from PDF")
                
            profile = {
                "Age": request.form.get("age"),
                "Grade Level": request.form.get("grade_level"),
                "Learning Goals": request.form.get("learning_goals"),
                "Subject Preferences": request.form.getlist("subject_preferences"),
                "Progress Data": request.form.get("progress_data"),
            }
            
            if not all([profile["Age"], profile["Grade Level"], profile["Learning Goals"]]):
                flash("Please fill in all required fields", "warning")
                return redirect(url_for("index"))

            # Store profile and PDF content
            session["profile"] = profile
            session["pdf_content"] = pdf_content
            
            # Check if this is an update or new profile creation
            is_update = session.get("profile_created", False)
            session["profile_created"] = True
            
            if is_update:
                flash("Profile updated successfully! You can now generate new career guidance or explore other features.", "success")
                return redirect(url_for("dashboard"))
            else:
                flash("Profile created successfully! Generating your career guidance...", "success")
                return redirect(url_for("career_guidance"))
                
        except Exception as e:
            logger.error(f"Error in setup: {e}", exc_info=True)
            flash("An error occurred during setup. Please try again.", "danger")
            return redirect(url_for("index"))

    return render_template("index.html")

@app.route("/career_guidance", methods=["GET", "POST"])
def career_guidance():
    if "profile" not in session:
        flash("Please complete your profile setup first.", "warning")
        return redirect(url_for("index"))

    try:
        if request.method == "POST":
            profile = session["profile"]
            pdf_content = session.get("pdf_content", "")
            
            if not pdf_content:
                flash("No resume content found. Please upload your resume again.", "warning")
                return redirect(url_for("index"))
            
            profile_str = '\n'.join(f"{k}: {v}" for k, v in profile.items())
            
            # Add logging
            logger.info(f"Generating career advice with profile length: {len(profile_str)} and PDF content length: {len(pdf_content)}")
            
            career_advice = generate_career_guidance(profile_str, pdf_content)
            
            # Check if error message
            if career_advice.startswith("Error"):
                flash(career_advice, "danger")
                return redirect(url_for("index"))
            
            session["career_advice_markdown"] = career_advice
            session["career_advice"] = markdown(career_advice)
            
            # Add logging
            logger.info("Career advice generated successfully")
            
            # Add success message for dashboard
            flash("Career guidance generated successfully! You can now download your report or explore other features.", "success")
            
        # Always render the career guidance page with current data
        return render_template(
            "career_guidance.html", 
            advice=session.get("career_advice", ""),
            profile=session.get("profile", {})
        )
                             
    except Exception as e:
        logger.error(f"Error in career guidance route: {e}", exc_info=True)
        flash("An error occurred generating career guidance. Please try again.", "danger")
        return redirect(url_for("index"))

@app.route("/skill_gap", methods=["GET", "POST"])
def skill_gap_analysis():
    if "user" not in session:
        return redirect(url_for("login"))
    
    try:    
        if request.method == "POST":
            target_role = request.form.get("target_role", "").strip()
            current_skills = request.form.getlist("current_skills")
            
            # Clean and validate current skills
            current_skills = [skill.strip() for skill in current_skills if skill.strip()]
            
            logger.info(f"Form data received - Target role: '{target_role}', Current skills: {current_skills}")
            
            if not target_role:
                flash("Please enter a target role", "warning")
                return render_template("skill_gap.html", show_results=False)
            
            logger.info(f"Processing skill gap analysis for role: {target_role}, skills: {current_skills}")
            
            # Generate skill gap analysis
            analysis = generate_skill_gap_analysis(target_role, current_skills, session.get("pdf_content", ""))
            
            # Validate analysis
            if not analysis or not isinstance(analysis, dict):
                logger.error(f"Invalid analysis response: {analysis}")
                flash("Error generating analysis. Please try again.", "danger")
                return render_template("skill_gap.html", show_results=False)
            
            # Log the analysis results for debugging
            logger.info(f"Analysis results:")
            logger.info(f"  Required skills: {analysis.get('required_skills', [])}")
            logger.info(f"  Missing skills: {analysis.get('missing_skills', [])}")
            logger.info(f"  Development priorities: {analysis.get('development_priorities', [])}")
            
            # Ensure missing_skills is not empty if there should be missing skills
            if not analysis.get('missing_skills') and analysis.get('required_skills'):
                # Manual calculation as fallback
                required = analysis['required_skills']
                current = current_skills
                missing = []
                
                for req_skill in required:
                    found = False
                    for curr_skill in current:
                        if req_skill.lower() in curr_skill.lower() or curr_skill.lower() in req_skill.lower():
                            found = True
                            break
                    if not found:
                        missing.append(req_skill)
                
                analysis['missing_skills'] = missing
                logger.info(f"Manually calculated missing skills: {missing}")
            
            # Get learning resources only for missing skills
            resources = {}
            if analysis.get('missing_skills'):
                resources = recommend_learning_resources(analysis['missing_skills'])
            
            # Store in session for potential later use
            session['skill_gap_analysis'] = {
                'analysis': analysis,
                'resources': resources,
                'target_role': target_role,
                'current_skills': current_skills
            }
            
            logger.info("Successfully completed skill gap analysis")
            
            return render_template(
                "skill_gap.html",
                analysis=analysis,
                resources=resources,
                target_role=target_role,
                current_skills=current_skills,
                show_results=True
            )
            
        return render_template("skill_gap.html", show_results=False)
        
    except Exception as e:
        logger.error(f"Error in skill gap analysis route: {e}", exc_info=True)
        flash("An error occurred during analysis. Please try again.", "danger")
        return render_template("skill_gap.html", show_results=False)
    
def get_suggested_roles(profile):
    """Generate suggested roles based on profile data"""
    suggested_roles = []
    
    # Add roles based on subject preferences
    subject_preferences = profile.get("Subject_Preferences", [])
    if isinstance(subject_preferences, str):
        subject_preferences = [subject_preferences]
        
    subject_role_map = {
        "Computer Science": ["Software Engineer", "Data Scientist", "Full Stack Developer", "DevOps Engineer"],
        "Mathematics": ["Data Analyst", "Quantitative Analyst", "Research Scientist", "Statistician"],
        "Science": ["Research Associate", "Lab Technician", "Scientific Writer", "Bioinformatics Specialist"],
        "Art": ["UX Designer", "Graphic Designer", "Art Director", "Creative Director"],
        "Business": ["Business Analyst", "Project Manager", "Product Manager", "Marketing Specialist"],
        "Engineering": ["Mechanical Engineer", "Electrical Engineer", "Systems Engineer", "Quality Assurance Engineer"],
        # Add more mappings as needed
    }
    
    # Get age and grade level for more targeted suggestions
    age = profile.get("Age", "")
    grade_level = profile.get("Grade_Level", "")
    
    # Add roles based on subject preferences
    for subject in subject_preferences:
        if subject in subject_role_map:
            suggested_roles.extend(subject_role_map[subject])
    
    # Add roles based on learning goals if available
    learning_goals = profile.get("Learning_Goals", "")
    if learning_goals:
        if "data" in learning_goals.lower():
            suggested_roles.extend(["Data Analyst", "Data Engineer", "Data Scientist"])
        if "ai" in learning_goals.lower() or "machine learning" in learning_goals.lower():
            suggested_roles.extend(["Machine Learning Engineer", "AI Researcher", "ML Ops Engineer"])
        if "web" in learning_goals.lower():
            suggested_roles.extend(["Web Developer", "Frontend Engineer", "UI Developer"])
        if "mobile" in learning_goals.lower():
            suggested_roles.extend(["Mobile App Developer", "iOS Developer", "Android Developer"])
        if "leadership" in learning_goals.lower() or "management" in learning_goals.lower():
            suggested_roles.extend(["Team Lead", "Product Manager", "Project Manager"])
    
    # Return unique roles, up to 6
    unique_roles = list(set(suggested_roles))
    return unique_roles[:6] if unique_roles else ["Software Engineer", "Data Analyst", "Project Manager"]

@app.route("/interview_prep", methods=["GET", "POST"])
def interview_preparation():
    if "user" not in session:
        return redirect(url_for("login"))
    
    # # Check if setup is required
    # if check_setup_required():
    #     flash("Please complete your profile setup first", "warning")
    #     return redirect(url_for("index"))
    
    try:
        if request.method == "POST":
            role = request.form.get("role")
            experience_level = request.form.get("experience_level")
            
            if not role or not experience_level:
                flash("Please fill in all required fields", "warning")
                return render_template("interview_prep.html", show_results=False)
            
            # Get profile and resume data from session
            profile = session.get("profile", {})
            resume_content = session.get("pdf_content", "")
            
            if not resume_content:
                flash("No resume content found. Please complete your profile setup first.", "warning")
                return redirect(url_for("index"))
            
            # Log the request
            logger.info(f"Processing interview prep for role: {role}, level: {experience_level}")
            
            try:
                # Use a try-except block specifically for the LLM call
                interview_data = generate_interview_prep(
                    role,
                    experience_level,
                    profile,
                    resume_content
                )
                
                # Validate interview_data
                if not interview_data or not isinstance(interview_data, dict):
                    logger.error(f"Invalid interview data format returned: {type(interview_data)}")
                    flash("There was an issue generating the interview preparation. Using backup content instead.", "warning")
                    
                    # Create a fallback structure if the return value isn't valid
                    interview_data = {
                        "technical_questions": [
                            {"question": f"What skills are most important for a {role} position?", 
                             "answer": "For this position, key skills typically include..."}
                        ],
                        "behavioral_questions": [
                            {"question": "Tell me about a challenging situation you faced at work.", 
                             "framework": "Use the STAR method: Situation, Task, Action, Result."}
                        ],
                        "questions_to_ask": [
                            "What does success look like in this role?",
                            "How would you describe the team culture?"
                        ]
                    }
                
                # For any missing keys, provide default values
                required_keys = ['technical_questions', 'behavioral_questions', 'questions_to_ask']
                for key in required_keys:
                    if key not in interview_data:
                        if key == 'technical_questions' or key == 'behavioral_questions':
                            interview_data[key] = [{"question": f"Sample {key.replace('_', ' ')}", 
                                                   "answer" if key == 'technical_questions' else "framework": "Sample response"}]
                        else:
                            interview_data[key] = ["Sample question to ask"]
                
                # Store in session
                session["interview_data"] = interview_data
                session["interview_role"] = role
                session["interview_level"] = experience_level
                
                logger.info("Successfully generated interview preparation data")
                
                # Pass the data to the template
                return render_template(
                    "interview_prep.html",
                    interview_data=interview_data,
                    role=role,
                    experience_level=experience_level,
                    show_results=True
                )
                
            except Exception as inner_e:
                # Handle exceptions specifically from the LLM call
                logger.error(f"Error during interview prep generation: {inner_e}", exc_info=True)
                flash("There was an issue connecting to the AI service. Please try again later.", "danger")
                return render_template("interview_prep.html", show_results=False)
                
        # For GET requests, simply render the template with form
        return render_template(
            "interview_prep.html", 
            show_results=False,
            profile=session.get("profile", {}),
            suggested_roles=get_suggested_roles(session.get("profile", {}))
        )
        
    except Exception as e:
        logger.error(f"Error in interview preparation route: {e}", exc_info=True)
        flash("An unexpected error occurred. Please try again.", "danger")
        return render_template("interview_prep.html", show_results=False)

@app.route("/career_comparison", methods=["GET", "POST"])
def career_comparison():
    if "user" not in session:
        return redirect(url_for("login"))
    
    try:
        if request.method == "POST":
            paths = request.form.getlist("career_paths")
            education_costs = request.form.getlist("education_costs")
            timeframes = request.form.getlist("timeframes")
            
            if not paths or not education_costs or not timeframes:
                flash("Please fill in all required fields", "warning")
                return redirect(url_for("career_comparison"))
            
            # Generate comparison data
            comparison_data = generate_career_comparison(
                paths,
                education_costs,
                timeframes
            )
            
            # Store in session
            session["comparison_data"] = comparison_data
            
            return render_template(
                "career_comparison.html",
                comparison_data=comparison_data,
                paths=paths,
                show_results=True
            )
            
        return render_template("career_comparison.html", show_results=False)
        
    except Exception as e:
        logger.error(f"Error in career comparison: {e}")
        flash("An error occurred. Please try again.", "danger")
        return render_template("career_comparison.html", show_results=False)

@app.route("/download_interview_prep")
def download_interview_prep():
    """Download interview preparation guide as DOCX"""
    if "interview_data" not in session:
        flash("No interview preparation data available.", "warning")
        return redirect(url_for("interview_preparation"))
        
    try:
        doc = Document()
        doc.add_heading("Interview Preparation Guide", 0)
        
        # Add technical questions
        doc.add_heading("Technical Questions", 1)
        for item in session["interview_data"]["technical_questions"]:
            doc.add_paragraph(f"Q: {item['question']}", style='Intense Quote')
            doc.add_paragraph(f"A: {item['answer']}")
            doc.add_paragraph()
            
        # Add behavioral questions
        doc.add_heading("Behavioral Questions", 1)
        for item in session["interview_data"]["behavioral_questions"]:
            doc.add_paragraph(f"Q: {item['question']}", style='Intense Quote')
            doc.add_paragraph(f"Framework: {item['framework']}")
            doc.add_paragraph()
            
        # Add questions to ask
        doc.add_heading("Questions to Ask the Interviewer", 1)
        for question in session["interview_data"]["questions_to_ask"]:
            doc.add_paragraph(question, style='List Bullet')
            
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            as_attachment=True,
            download_name="interview_prep_guide.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        logger.error(f"Error creating interview prep document: {e}")
        flash("Error creating document. Please try again.", "danger")
        return redirect(url_for("interview_preparation"))

@app.route("/download_doc")
def download_doc():
    """Generate and download the career guidance as a .docx file"""
    if "career_advice_markdown" not in session:
        flash("No career advice available to download.", "warning")
        return redirect(url_for("career_guidance"))

    try:
        doc_io = create_doc_file(session["career_advice_markdown"])
        return send_file(
            doc_io,
            as_attachment=True,
            download_name="career_guidance.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        logger.error(f"Error downloading document: {e}")
        flash("An error occurred creating the document. Please try again.", "danger")
        return redirect(url_for("career_guidance"))

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

@app.route("/api/save_progress", methods=["POST"])
def save_progress():
    """API endpoint to save user progress"""
    if "user" not in session:
        return jsonify({"error": "Not authenticated"}), 401
        
    try:
        data = request.get_json()
        # Save progress data to session
        if "progress" not in session:
            session["progress"] = {}
        
        session["progress"][data["section"]] = {
            "completed": data["completed"],
            "timestamp": datetime.now().isoformat()
        }
        
        return jsonify({"success": True})
    except Exception as e:
        logger.error(f"Error saving progress: {e}")
        return jsonify({"error": "Failed to save progress"}), 500

@app.route("/api/get_progress")
def get_progress():
    """API endpoint to get user progress"""
    if "user" not in session:
        return jsonify({"error": "Not authenticated"}), 401
        
    try:
        progress = session.get("progress", {})
        return jsonify(progress)
    except Exception as e:
        logger.error(f"Error getting progress: {e}")
        return jsonify({"error": "Failed to get progress"}), 500

@app.errorhandler(Exception)
def handle_error(e):
    logger.error(f"Unhandled error: {e}")
    flash("An unexpected error occurred. Please try again.", "danger")
    return redirect(url_for("index"))
def open_browser():
    """Open browser after a short delay"""
    time.sleep(1)  # Give the server a second to start
    webbrowser.open('http://127.0.0.1:5000/')

if __name__ == "__main__":
    threading.Thread(target=open_browser).start()
    app.run(debug=True)